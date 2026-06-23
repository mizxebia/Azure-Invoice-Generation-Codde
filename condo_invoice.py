import json
import logging
import os
import argparse
import time
import traceback
from datetime import datetime
from pathlib import Path

import pandas as pd
import requests
from msal import ConfidentialClientApplication
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn


CONFIG_PATH = Path(__file__).with_name("config.json")
GRAPH_BASE = "https://graph.microsoft.com/v1.0"
CLOSING_PDF_FILE_COLUMN = "cr109_closingticketdetailspdf"

ENVIRONMENT_SETTINGS = {
    "DEV": {
        "user_email_key": "user_email_dev",
        "user_email": "akambotdev1@akam.com",
        "condo_template_path": (
            "New Sales RPA/DEV/Excel Sheets/CondoInvoiceTemplate.docx"
        ),
        "base_folder": "New Sales RPA/DEV/BotShareDrive/InProgress",
    },
    "UAT": {
        "user_email_key": "user_email_uat",
        "user_email": "akambotuat2@akam.com",
        "condo_template_path": (
            "New Sales RPA/UAT/Excel Sheets/CondoInvoiceTemplate.docx"
        ),
        "base_folder": "New Sales RPA/UAT/BotShareDrive/InProgress",
    },
    "PROD": {
        "user_email_key": "user_email_prod",
        "user_email": "akambotnewsalesclosure@akam.com",
        "condo_template_path": (
            "New Sales RPA/PROD/Excel Sheets/CondoInvoiceTemplate.docx"
        ),
        "base_folder": "New Sales RPA/PROD/BotShareDrive/InProgress",
    },
}


PAYABLE_MAP = {
    716070000: "Building",
    716070001: "AKAM",
    716070002: "Other",
}

TRANSACTION_TYPE_DEAL_MAP = {
    396620000: "All Cash",
    396620001: "Financing",
    396620002: "Transfer",
    396620003: "Trust Transfer",
}

DUE_AT_CLOSING_MAP = {
    396620000: "Adjournment Fee",
    396620001: "Admin Fee",
    396620002: "Air Conditioning Fee",
    396620003: "AKAM Processing Fee",
    396620004: "Appliance Fee",
    396620005: "Application Fee",
    396620006: "Arrears",
    396620007: "Assessment",
    396620008: "Assignment of Share",
    396620009: "Background Check",
    396620010: "Building Admin Fee",
    396620011: "Cable Charges",
    396620012: "Capital Assessment Fee",
    396620013: "Carpet Deposit",
    396620014: "Change of Occupancy",
    396620015: "Closing Fee (Non-Refundable)",
    396620016: "Contribution Fee (Non-Refundable)",
    396620017: "Contribution Reserves",
    396620018: "COOP Prospectus",
    396620019: "COOP Questionnaire",
    396620020: "Credit Report / Check",
    396620021: "Electric Fee",
    396620022: "Elevator Fee",
    396620023: "Energy Charge",
    396620024: "Escrow Maintenance",
    396620025: "Estate Review Fee",
    396620026: "Expediting Fee",
    396620027: "Flip Tax",
    396620028: "Guarantee Fee",
    396620029: "Inspection",
    396620030: "Legal Fee",
    396620031: "Lost Stock & Lease",
    396620032: "Maintenance Fees",
    396620033: "Major/Minor Alteration Fee",
    396620034: "Messenger",
    396620035: "Meter Fee",
    396620036: "Mortgage Questionnaire",
    396620037: "Move In/Out Deposit",
    396620038: "Move In/Out Fee",
    396620039: "Other",
    396620040: "Over-Time Fee",
    396620041: "Parking",
    396620042: "POA Fee",
    396620043: "Processing Fee",
    396620044: "Purchaser Fee (Transfer Fee)",
    396620045: "Real Estate Tax",
    396620046: "Recognition Agreement",
    396620047: "Repair Charge",
    396620048: "Resident Manager Contribution",
    396620049: "Security Deposit",
    396620050: "Service Fee",
    396620051: "Stock Transfer Fee",
    396620052: "Storage Unit",
    396620053: "Sublet Deposit",
    396620054: "Sublet Fee",
    396620055: "Transfer Fee",
    396620056: "Utilities",
    396620057: "Waiver Fee",
    396620058: "Working Capital",
}


def load_config():
    with open(CONFIG_PATH, "r", encoding="utf-8") as config_file:
        return json.load(config_file)


def parse_args():
    parser = argparse.ArgumentParser(
        description="Generate condo invoice Word documents."
    )
    parser.add_argument(
        "environment",
        nargs="?",
        help="Target environment: DEV, UAT, or PROD.",
    )
    parser.add_argument(
        "--env",
        dest="env",
        help="Target environment: DEV, UAT, or PROD.",
    )
    return parser.parse_args()


def normalize_environment(env_value):
    env = (env_value or "DEV").strip().upper()

    if env not in ENVIRONMENT_SETTINGS:
        valid_envs = ", ".join(ENVIRONMENT_SETTINGS)
        raise ValueError(f"Invalid environment '{env}'. Use one of: {valid_envs}.")

    return env


def get_runtime_settings(config, env):
    storage = config["storage"]
    auth_config = config["auth"]
    dv = config["dataverse"]
    env_settings = ENVIRONMENT_SETTINGS[env]
    dataverse = auth_config.get("dataverse_by_env", {}).get(env, {})

    tables = (
        dv.get("tables_by_env", {}).get(env)
        or dv.get(f"tables_{env.lower()}")
        or dv["tables"]
    )

    return {
        "env": env,
        "user_email": storage.get(
            env_settings["user_email_key"],
            env_settings["user_email"],
        ),
        "condo_template_path": env_settings["condo_template_path"],
        "base_folder": env_settings["base_folder"],
        "dataverse_tables": tables,
        "dataverse_url": dataverse.get(
            "dataverse_url",
            auth_config["dataverse_url"],
        ),
        "dataverse_scope": dataverse.get(
            "dataverse_scope",
            auth_config["dataverse_scope"],
        ),
    }


config = load_config()
auth = config["auth"]

logging.basicConfig(
    filename=config["logging"]["log_file"],
    level=getattr(logging, config["logging"]["log_level"]),
    format="%(asctime)s | %(levelname)s | %(message)s",
)
logger = logging.getLogger(__name__)


def get_access_token(scope):
    try:
        authority_url = f"https://login.microsoftonline.com/{auth['tenant_id']}"

        app = ConfidentialClientApplication(
            auth["client_id"],
            authority=authority_url,
            client_credential=auth["client_secret"],
        )

        token_response = app.acquire_token_for_client(scopes=[scope])

        if "access_token" not in token_response:
            raise Exception(token_response)

        return token_response["access_token"]

    except Exception as error:
        logger.error(traceback.format_exc())
        print(error)
        return None


def fetch_table_from_dataverse_url(
    dataverse_url,
    table_name,
    token,
    ticket_column,
    ticket_value,
):
    headers = {
        "Authorization": f"Bearer {token}",
        "Accept": "application/json",
    }

    url = (
        f"{dataverse_url}/api/data/v9.2/{table_name}"
        f"?$filter={ticket_column} eq '{ticket_value}'"
    )

    all_records = []

    while url:
        response = requests.get(url, headers=headers)

        if response.status_code != 200:
            raise Exception(response.text)

        data = response.json()
        all_records.extend(data.get("value", []))
        url = data.get("@odata.nextLink")

    return all_records


def validate_required_dataverse_data(
    env,
    ticket_value,
    closing_data,
    invoice_data,
    closing_table,
    invoice_table,
):
    errors = []

    if not closing_data:
        errors.append(
            "No Closing Ticket Details records found "
            f"in table '{closing_table}' for ticket '{ticket_value}'."
        )

    if not invoice_data:
        errors.append(
            "No Invoice Details records found "
            f"in table '{invoice_table}' for ticket '{ticket_value}'."
        )

    if errors:
        message = (
            f"Dataverse data missing for environment '{env}'. "
            + " ".join(errors)
        )
        raise ValueError(message)


def get_choice_label(mapping, value):
    if value in (None, "") or pd.isna(value):
        return ""

    try:
        normalized_value = int(value)
    except (TypeError, ValueError):
        normalized_value = value

    return mapping.get(normalized_value, "")


def get_row_id(row, table_name):
    primary_key = table_name[:-2] + "id" if table_name.endswith("es") else f"{table_name}id"

    if primary_key in row:
        return row[primary_key]

    for key, value in row.items():
        if key.endswith("id") and not key.startswith("_"):
            return value

    raise ValueError(f"Could not determine Dataverse row id for table '{table_name}'.")


def _headers(token, extra=None):
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json",
    }

    if extra:
        headers.update(extra)

    return headers


def _ensure_folder(token, user_email, folder_path):
    check_url = (
        f"{GRAPH_BASE}/users/{user_email}"
        f"/drive/root:/{folder_path}"
    )

    response = requests.get(check_url, headers=_headers(token))

    if response.status_code == 200:
        return

    parts = folder_path.rsplit("/", 1)
    parent_path = parts[0] if len(parts) == 2 else ""
    folder_name = parts[-1]

    if parent_path:
        create_url = (
            f"{GRAPH_BASE}/users/{user_email}"
            f"/drive/root:/{parent_path}:/children"
        )
    else:
        create_url = f"{GRAPH_BASE}/users/{user_email}/drive/root/children"

    body = {
        "name": folder_name,
        "folder": {},
        "@microsoft.graph.conflictBehavior": "replace",
    }

    response = requests.post(create_url, headers=_headers(token), json=body)

    if response.status_code not in [200, 201]:
        raise Exception(f"Failed to create folder '{folder_path}': {response.text}")

    print(f"\nFolder created: {folder_path}")


def _get_file_metadata(token, user_email, file_path):
    url = (
        f"{GRAPH_BASE}/users/{user_email}"
        f"/drive/root:/{file_path}"
    )

    response = requests.get(url, headers=_headers(token))

    if response.status_code == 200:
        return response.json()

    if response.status_code == 404:
        return None

    raise Exception(f"Error checking file '{file_path}': {response.text}")


def _move_and_rename_file(
    token,
    user_email,
    source_path,
    destination_folder_path,
    new_name,
):
    folder_url = (
        f"{GRAPH_BASE}/users/{user_email}"
        f"/drive/root:/{destination_folder_path}"
    )
    folder_response = requests.get(folder_url, headers=_headers(token))

    if folder_response.status_code != 200:
        raise Exception(
            f"Cannot find destination folder '{destination_folder_path}': "
            f"{folder_response.text}"
        )

    folder_id = folder_response.json()["id"]
    move_url = (
        f"{GRAPH_BASE}/users/{user_email}"
        f"/drive/root:/{source_path}"
    )
    body = {
        "parentReference": {"id": folder_id},
        "name": new_name,
    }

    response = requests.patch(move_url, headers=_headers(token), json=body)

    if response.status_code != 200:
        raise Exception(f"Failed to move/rename file: {response.text}")

    print(f"\nArchived: {new_name}")


def setup_ticket_folders(token, user_email, base_folder, ticket_id):
    ticket_folder = f"{base_folder}/{ticket_id}"
    active_folder = f"{ticket_folder}/Active"
    inactive_folder = f"{ticket_folder}/Inactive"

    _ensure_folder(token, user_email, ticket_folder)
    _ensure_folder(token, user_email, active_folder)
    _ensure_folder(token, user_email, inactive_folder)

    print(f"\n📁 Folder Structure Created:")
    print(f"   OneDrive User: {user_email}")
    print(f"   Base Path: {base_folder}")
    print(f"   Ticket Folder: {ticket_folder}")
    print(f"   Active Folder: {active_folder}")
    print(f"   Inactive Folder: {inactive_folder}")

    return active_folder, inactive_folder


def archive_active_file_if_exists(
    token,
    user_email,
    active_folder,
    inactive_folder,
    output_file_name,
):
    active_file_path = f"{active_folder}/{output_file_name}"
    existing = _get_file_metadata(token, user_email, active_file_path)

    if existing is None:
        print("\nNo existing file in Active - skipping archive")
        return

    now = datetime.now()
    date_str = now.strftime("%Y-%m-%d")
    datetime_str = now.strftime("%Y-%m-%d_%H-%M-%S")
    name_parts = output_file_name.rsplit(".", 1)

    if len(name_parts) == 2:
        archived_name = f"{name_parts[0]}_{datetime_str}.{name_parts[1]}"
    else:
        archived_name = f"{output_file_name}_{datetime_str}"

    date_folder = f"{inactive_folder}/{date_str}"
    _ensure_folder(token, user_email, date_folder)
    _move_and_rename_file(
        token,
        user_email,
        active_file_path,
        date_folder,
        archived_name,
    )

    print(f"\nArchived to: Inactive/{date_str}/{archived_name}")


def archive_closing_form_files_if_exist(
    token,
    user_email,
    active_folder,
    inactive_folder,
    ticket_id,
):
    """Archive any Closing_Form files (pdf/xlsx) found in the active folder.

    Looks for files whose names start with '<ticket_id>_Closing_Form' and
    moves each one to Inactive/<today's date>/ with a timestamp suffix.
    """
    list_url = (
        f"{GRAPH_BASE}/users/{user_email}"
        f"/drive/root:/{active_folder}:/children"
        f"?$select=name,id"
    )

    response = requests.get(list_url, headers=_headers(token))

    if response.status_code == 404:
        # Active folder doesn't exist yet — nothing to archive
        return

    if response.status_code != 200:
        raise Exception(
            f"Failed to list Active folder contents: {response.text}"
        )

    items = response.json().get("value", [])
    prefix = f"{ticket_id}_Closing_Form"

    matches = [
        item for item in items
        if item["name"].startswith(prefix)
        and item["name"].lower().endswith((".pdf", ".xlsx"))
    ]

    if not matches:
        print(f"\nNo Closing_Form files found in Active — skipping archive")
        return

    now = datetime.now()
    date_str = now.strftime("%Y-%m-%d")
    datetime_str = now.strftime("%Y-%m-%d_%H-%M-%S")
    date_folder = f"{inactive_folder}/{date_str}"
    _ensure_folder(token, user_email, date_folder)

    for item in matches:
        original_name = item["name"]
        name_parts = original_name.rsplit(".", 1)
        if len(name_parts) == 2:
            archived_name = f"{name_parts[0]}_{datetime_str}.{name_parts[1]}"
        else:
            archived_name = f"{original_name}_{datetime_str}"

        source_path = f"{active_folder}/{original_name}"
        _move_and_rename_file(
            token,
            user_email,
            source_path,
            date_folder,
            archived_name,
        )
        print(f"\nArchived Closing_Form file to: Inactive/{date_str}/{archived_name}")


def download_template_from_onedrive(token, user_email, template_path, local_path):
    url = (
        f"{GRAPH_BASE}/users/{user_email}"
        f"/drive/root:/{template_path}:/content"
    )

    response = requests.get(url, headers=_headers(token))

    if response.status_code != 200:
        raise Exception(f"Failed to download template: {response.text}")

    with open(local_path, "wb") as file:
        file.write(response.content)

    print(f"\nTemplate downloaded: {local_path}")


def upload_file_to_onedrive(
    token,
    user_email,
    folder_path,
    local_file_path,
    onedrive_file_name,
):
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    }

    url = (
        f"{GRAPH_BASE}/users/{user_email}"
        f"/drive/root:/{folder_path}/"
        f"{onedrive_file_name}:/content"
    )

    with open(local_file_path, "rb") as file:
        file_content = file.read()

    response = requests.put(url, headers=headers, data=file_content)

    if response.status_code not in [200, 201]:
        raise Exception(response.text)

    full_path = f"{folder_path}/{onedrive_file_name}"
    print(f"\n✓ File uploaded successfully!")
    print(f"  OneDrive User: {user_email}")
    print(f"  OneDrive Path: {full_path}")


def upload_file_to_dataverse_file_column(
    dataverse_url,
    table_name,
    row_id,
    token,
    file_column,
    local_file_path,
    file_name,
):
    url = (
        f"{dataverse_url}/api/data/v9.2/{table_name}"
        f"({row_id})/{file_column}"
    )
    headers = {
        "Authorization": f"Bearer {token}",
        "OData-MaxVersion": "4.0",
        "OData-Version": "4.0",
        "If-None-Match": "null",
        "Accept": "application/json",
        "Content-Type": "application/octet-stream",
        "x-ms-file-name": file_name,
    }

    with open(local_file_path, "rb") as file:
        response = requests.patch(url, headers=headers, data=file.read())

    if response.status_code != 204:
        raise Exception(
            f"Failed to upload PDF to Dataverse file column "
            f"'{file_column}': {response.text}"
        )

    print(f"Uploaded PDF to Dataverse column: {table_name}.{file_column}")


def get_content_control_by_tag(element, tag):
    """Find content control by tag name in XML"""
    namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    # Search for all sdt (structured document tags / content controls)
    for sdt in element.findall('.//w:sdt', namespace):
        # Get the tag element
        tag_elem = sdt.find('.//w:tag', namespace)
        if tag_elem is not None:
            tag_val = tag_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            if tag_val == tag:
                return sdt
    return None


def get_content_control_by_tag_anywhere(doc, tag):
    """Find content control by tag name searching body, headers, and footers"""
    # Search main body
    sdt = get_content_control_by_tag(doc._element, tag)
    if sdt is not None:
        return sdt

    # Search headers and footers
    for section in doc.sections:
        for part in [
            section.header,
            section.footer,
            section.even_page_header,
            section.even_page_footer,
            section.first_page_header,
            section.first_page_footer,
        ]:
            try:
                if part and part._element is not None:
                    sdt = get_content_control_by_tag(part._element, tag)
                    if sdt is not None:
                        return sdt
            except Exception:
                pass

    return None


def create_table_in_content_control(sdt, invoice_df, building_address):
    """Create a table inside a content control for invoice data"""
    from docx.table import _Cell
    from docx.oxml import OxmlElement
    
    namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    if len(invoice_df) == 0:
        return set_content_control_text(sdt, "No charges", None, None, 'auto')
    
    # Find the sdtContent element
    content = sdt.find('.//w:sdtContent', namespace)
    if content is None:
        return False
    
    # Clear existing content
    for child in list(content):
        content.remove(child)
    
    # Create a table with rows for each invoice item
    # Table structure: Number | Amount | Description | Payable To
    num_rows = len(invoice_df)
    
    # Create table element
    tbl = parse_xml(f'<w:tbl xmlns:w="{namespace["w"]}"></w:tbl>')
    
    # Table properties
    tblPr = parse_xml(f'''<w:tblPr xmlns:w="{namespace['w']}">
        <w:tblW w:w="5000" w:type="pct"/>
        <w:tblBorders>
            <w:top w:val="none" w:sz="0"/>
            <w:left w:val="none" w:sz="0"/>
            <w:bottom w:val="none" w:sz="0"/>
            <w:right w:val="none" w:sz="0"/>
            <w:insideH w:val="none" w:sz="0"/>
            <w:insideV w:val="none" w:sz="0"/>
        </w:tblBorders>
    </w:tblPr>''')
    tbl.append(tblPr)
    
    # Table grid (4 columns)
    tblGrid = parse_xml(f'''<w:tblGrid xmlns:w="{namespace['w']}">
        <w:gridCol w:w="500"/>
        <w:gridCol w:w="4500"/>
    </w:tblGrid>''')
    tbl.append(tblGrid)
    
    # Add rows for each invoice
    for row_number, (idx, row_data) in enumerate(invoice_df.iterrows(), start=1):
        amount = row_data.get("cr7de_amount", "")
        description = get_choice_label(
            DUE_AT_CLOSING_MAP,
            row_data.get("cr109_dueatclosing", ""),
        )
        payable_to_code = row_data.get("cr7de_payableto", "")
        payable_to = get_choice_label(PAYABLE_MAP, payable_to_code)

        if payable_to == "Building":
            if building_address and building_address.strip():
                payable_to = building_address
            else:
                payable_to = "Building"
        elif payable_to == "AKAM":
            payable_to = "AKAM Associates, Inc"
        elif payable_to == "Other":
            payable_to = row_data.get("cr109_otherpayableto", "") or "Other"

        # Format amount
        try:
            amount_float = float(amount) if amount else 0.0
            amount_str = f"${amount_float:,.2f}"
        except (TypeError, ValueError):
            amount_str = f"${amount}"
        
        # Create row
        tr = parse_xml(f'<w:tr xmlns:w="{namespace["w"]}"></w:tr>')
        
        # Cell 1: Number (sequential starting from 1)
        tc1 = parse_xml(f'''<w:tc xmlns:w="{namespace['w']}">
            <w:tcPr><w:tcW w:w="500" w:type="dxa"/></w:tcPr>
            <w:p>
                <w:r>
                    <w:rPr><w:color w:val="auto"/></w:rPr>
                    <w:t>{row_number})</w:t>
                </w:r>
            </w:p>
        </w:tc>''')
        tr.append(tc1)
        
        # Cell 2: Amount - Description made payable to Payable To
        detail_text = f"{amount_str} - {description} made payable to {payable_to}"
        tc2 = parse_xml(f'''<w:tc xmlns:w="{namespace['w']}">
            <w:tcPr><w:tcW w:w="4500" w:type="dxa"/></w:tcPr>
            <w:p>
                <w:r>
                    <w:rPr><w:color w:val="auto"/></w:rPr>
                    <w:t xml:space="preserve">{detail_text}</w:t>
                </w:r>
            </w:p>
        </w:tc>''')
        tr.append(tc2)
        
        tbl.append(tr)
    
    content.append(tbl)
    return True


def set_content_control_text(sdt, text, font_name=None, font_size=None, color='auto'):
    """Set text in a content control with optional formatting.

    Handles three structural variants found in Word content controls:
      1. sdtContent > w:p  (block content control)
      2. sdtContent > w:r  (inline/run-level content control, no paragraph)
      3. sdtContent exists but is empty — a new paragraph+run is created
    """
    namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    W = namespace['w']

    # Find the sdtContent element
    content = sdt.find('.//w:sdtContent', namespace)
    if content is None:
        print(f"    [DEBUG] sdtContent not found in sdt")
        return False

    # ── helper: build a single run element ──────────────────────────────────
    def make_run(line_text):
        run = parse_xml(f'<w:r xmlns:w="{W}"></w:r>')
        rPr = parse_xml(f'<w:rPr xmlns:w="{W}"></w:rPr>')
        run.insert(0, rPr)

        if font_name:
            rPr.append(parse_xml(
                f'<w:rFonts xmlns:w="{W}" '
                f'w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>'
            ))
        if font_size:
            size_val = str(int(font_size) * 2)
            rPr.append(parse_xml(f'<w:sz xmlns:w="{W}" w:val="{size_val}"/>'))
            rPr.append(parse_xml(f'<w:szCs xmlns:w="{W}" w:val="{size_val}"/>'))
        if color in ('auto', 'black'):
            rPr.append(parse_xml(f'<w:color xmlns:w="{W}" w:val="auto"/>'))

        t_elem = parse_xml(f'<w:t xmlns:w="{W}" xml:space="preserve"></w:t>')
        t_elem.text = line_text
        run.append(t_elem)
        return run

    lines = str(text).split('\n')

    # ── Variant 1: block content control — has a <w:p> child ────────────────
    para = content.find('.//w:p', namespace)
    if para is not None:
        # Remove all existing runs (but preserve paragraph properties)
        for run in para.findall('w:r', namespace):
            para.remove(run)
        # Also remove any leftover w:ins / w:del that wrap runs
        for wrapper in para.findall('w:ins', namespace) + para.findall('w:del', namespace):
            para.remove(wrapper)

        for i, line in enumerate(lines):
            para.append(make_run(line))
            if i < len(lines) - 1:
                para.append(parse_xml(f'<w:r xmlns:w="{W}"><w:br/></w:r>'))
        return True

    # ── Variant 2: inline content control — sdtContent contains w:r directly ─
    for old_run in content.findall('w:r', namespace):
        content.remove(old_run)

    for i, line in enumerate(lines):
        content.append(make_run(line))
        if i < len(lines) - 1:
            content.append(parse_xml(f'<w:r xmlns:w="{W}"><w:br/></w:r>'))

    # If we got here without finding any runs either, the content was empty.
    # Return True anyway — we appended runs directly to sdtContent.
    return True


def populate_content_controls_by_tag(doc, tag_mappings, buyer_df=None, seller_df=None, building_address=None):
    """Populate content controls using their tag names"""
    namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    populated_tags = []
    
    # Get the document's XML element
    doc_element = doc._element
    
    for tag, value_info in tag_mappings.items():
        # Handle both simple string values and tuples with formatting
        if isinstance(value_info, tuple):
            if len(value_info) == 3:
                value, font_name, font_size = value_info
                color = 'auto'
            elif len(value_info) == 4:
                value, font_name, font_size, color = value_info
            else:
                value = value_info[0]
                font_name = font_size = None
                color = 'auto'
        else:
            value, font_name, font_size, color = value_info, None, None, 'auto'
        
        # Find the content control with this tag (searches body, headers, footers)
        sdt = get_content_control_by_tag_anywhere(doc, tag)
        
        if sdt is not None:
            # Special handling for invoice tables
            if tag == "BuyerCheques" and buyer_df is not None:
                success = create_table_in_content_control(sdt, buyer_df, building_address)
            elif tag == "SellerCheques" and seller_df is not None:
                success = create_table_in_content_control(sdt, seller_df, building_address)
            else:
                success = set_content_control_text(sdt, value, font_name, font_size, color)
            
            if success:
                populated_tags.append(tag)
                print(f"  ✓ Populated tag '{tag}'")
            else:
                print(f"  ✗ Failed to populate tag '{tag}'")
        else:
            print(f"  ⚠ Tag '{tag}' not found in document")
    
    return populated_tags


def replace_text_in_paragraph(paragraph, replacements):
    """Replace placeholder text in a paragraph while preserving formatting"""
    for key, value in replacements.items():
        if key in paragraph.text:
            # Handle replacements in runs to preserve formatting
            for run in paragraph.runs:
                if key in run.text:
                    run.text = run.text.replace(key, str(value))


def replace_text_in_document(doc, replacements):
    """Replace all placeholder text in the document"""
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)


def format_invoice_line(amount, description, payable_to):
    """Format a single invoice line item"""
    # Format amount with currency
    try:
        amount_float = float(amount) if amount else 0.0
        amount_str = f"${amount_float:,.2f}"
    except (TypeError, ValueError):
        amount_str = f"${amount}"
    
    # Build the line
    line = f"{amount_str} - {description} made payable to {payable_to}"
    return line


def build_invoice_table_html(invoice_df, building_address):
    """Build formatted invoice table from dataframe"""
    if len(invoice_df) == 0:
        return "No charges"
    
    lines = []
    
    for idx, row in invoice_df.iterrows():
        amount = row.get("cr7de_amount", "")
        description = get_choice_label(
            DUE_AT_CLOSING_MAP,
            row.get("cr109_dueatclosing", ""),
        )
        payable_to_code = row.get("cr7de_payableto", "")
        payable_to = get_choice_label(PAYABLE_MAP, payable_to_code)

        if payable_to == "Building":
            if building_address and building_address.strip():
                payable_to = building_address
            else:
                payable_to = "Building"
        elif payable_to == "AKAM":
            payable_to = "AKAM Associates, Inc"
        elif payable_to == "Other":
            payable_to = row.get("cr109_otherpayableto", "") or "Other"

        # Format amount with currency
        try:
            amount_float = float(amount) if amount else 0.0
            amount_str = f"${amount_float:,.2f}"
        except (TypeError, ValueError):
            amount_str = f"${amount}"
        
        # Build table row format: number) amount - description made payable to payable_to
        line = f"{idx + 1})\t{amount_str} - {description} made payable to {payable_to}"
        lines.append(line)
    
    return "\n".join(lines)


def populate_word_template(
    template_path,
    output_path,
    closing_ticket_df,
    invoice_df,
):
    """Populate Word template with data using Content Control tags"""
    doc = Document(template_path)
    
    # Get data from closing ticket
    row = closing_ticket_df.iloc[0]
    
    property_address = row.get("cr7de_buildingaddress", "")
    unit = row.get("cr7de_unitnumber", "")
    building_address = row.get("cr109_legalname", "") or row.get("cr7de_buildingaddress", "")

    # Normalize NaN/None to empty string
    if pd.isna(property_address) if not isinstance(property_address, str) else not property_address:
        property_address = ""
    if pd.isna(unit) if not isinstance(unit, str) else not unit:
        unit = ""
    if pd.isna(building_address) if not isinstance(building_address, str) else not building_address:
        building_address = ""

    print(f"\n  [DEBUG] cr7de_buildingaddress = '{property_address}'")
    print(f"  [DEBUG] cr7de_unitnumber     = '{unit}'")
    print(f"  [DEBUG] Available closing ticket columns: {list(closing_ticket_df.columns)}")
    closing_agent_name = row.get("cr7de_closingagentname", "")
    closing_agent_phone = row.get("cr7de_closingagentphone", "")
    closing_agent_email = row.get("cr7de_closingagentemail", "")
    
    # Separate buyer and seller invoices
    seller_df = invoice_df[invoice_df["cr7de_paidby"] == 716070000]
    buyer_df = invoice_df[invoice_df["cr7de_paidby"] == 716070001]
    
    print("\n📝 Populating Content Controls...")
    
    # Map tags to values with formatting
    # Format: tag: value OR tag: (value, font_name, font_size, color)
    # Note: BuyerCheques and SellerCheques will be populated as tables
    tag_mappings = {
        "Address": (property_address, None, None, 'auto'),
        "Unit": (unit, None, None, 'auto'),
        "BuyerCheques": ("", None, None, 'auto'),  # Will be replaced with table
        "SellerCheques": ("", None, None, 'auto'),  # Will be replaced with table
        "Agent Number": (closing_agent_phone, None, None, 'auto'),
        "Agent Email": (closing_agent_email, None, None, 'auto'),
        "ClosingAgentName": (closing_agent_name, None, None, 'auto'),
        "ClosingAgentSignature": (closing_agent_name, "Cochocib Script Latin Pro", 30, 'auto'),
    }
    
    # Populate all content controls by tag
    populated_tags = populate_content_controls_by_tag(
        doc, 
        tag_mappings,
        buyer_df=buyer_df,
        seller_df=seller_df,
        building_address=building_address
    )
    
    # Save the populated document
    doc.save(output_path)
    
    print(f"\n✓ Word document populated successfully!")
    print(f"  File: {output_path}")
    print(f"  Property: {property_address}")
    print(f"  Unit: {unit}")
    print(f"  Buyer fees: {len(buyer_df)} items")
    print(f"  Seller fees: {len(seller_df)} items")
    print(f"  Closing Agent: {closing_agent_name}")
    print(f"  Tags populated: {len(populated_tags)}/{len(tag_mappings)}")


def convert_word_to_pdf_onedrive(token, user_email, file_path, pdf_output_path):
    """Convert Word document to PDF using OneDrive API"""
    url = (
        f"{GRAPH_BASE}/users/{user_email}"
        f"/drive/root:/{file_path}:/content"
        "?format=pdf"
    )

    response = requests.get(url, headers=_headers(token))
    print("\nPDF conversion response:", response.status_code)

    if response.status_code != 200:
        print(response.text)
        raise Exception("PDF conversion failed")

    with open(pdf_output_path, "wb") as file:
        file.write(response.content)

    print("\nPDF generated successfully")


def main():
    args = parse_args()
    env = normalize_environment(
        args.env
        or args.environment
        or os.getenv("APP_ENV")
        or os.getenv("ENV")
        or config.get("environment")
    )
    runtime = get_runtime_settings(config, env)

    dv_token = get_access_token(runtime["dataverse_scope"])
    graph_token = get_access_token(config["auth"]["graph_scope"])

    dv = config["dataverse"]
    dv_tables = runtime["dataverse_tables"]
    dataverse_url = runtime["dataverse_url"]

    ticket_column = dv["columns"]["ticket_id"]
    ticket_value = dv["filter"]["ticket_id"]
    template_path = runtime["condo_template_path"]
    user_email = runtime["user_email"]
    base_folder = runtime["base_folder"]
    output_word_name = f"{ticket_value}_Condo_Invoice.docx"
    pdf_output_name = f"{ticket_value}_Condo_Invoice.pdf"
    
    local_template_path = f"temp_template_{ticket_value}.docx"
    local_output_path = f"temp_output_{ticket_value}.docx"
    local_pdf_path = pdf_output_name

    print(f"\nEnvironment: {env}")
    print(f"OneDrive user: {user_email}")
    print(f"Dataverse URL: {dataverse_url}")
    print(f"Template Path: {template_path}")
    print(f"Base Folder: {base_folder}")
    print(f"Ticket ID: {ticket_value}")

    closing_table = dv_tables["closing_ticket_details"]
    invoice_table = dv_tables["invoice_details"]

    # Fetch data from Dataverse
    closing_data = fetch_table_from_dataverse_url(
        dataverse_url,
        closing_table,
        dv_token,
        ticket_column,
        ticket_value,
    )
    invoice_data = fetch_table_from_dataverse_url(
        dataverse_url,
        invoice_table,
        dv_token,
        ticket_column,
        ticket_value,
    )

    print(f"Closing Ticket Details records: {len(closing_data)}")
    print(f"Invoice Details records: {len(invoice_data)}")

    validate_required_dataverse_data(
        env,
        ticket_value,
        closing_data,
        invoice_data,
        closing_table,
        invoice_table,
    )

    df_closing = pd.DataFrame(closing_data)
    df_invoice = pd.DataFrame(invoice_data)
    closing_row_id = get_row_id(closing_data[0], closing_table)

    # Setup folders
    active_folder, inactive_folder = setup_ticket_folders(
        graph_token,
        user_email,
        base_folder,
        ticket_value,
    )

    # Archive existing condo invoice file if present
    archive_active_file_if_exists(
        graph_token,
        user_email,
        active_folder,
        inactive_folder,
        output_word_name,
    )

    # Archive any Closing_Form files (pdf/xlsx) if present
    archive_closing_form_files_if_exist(
        graph_token,
        user_email,
        active_folder,
        inactive_folder,
        ticket_value,
    )

    # Download template
    download_template_from_onedrive(
        graph_token,
        user_email,
        template_path,
        local_template_path,
    )

    # Populate Word template
    populate_word_template(
        local_template_path,
        local_output_path,
        df_closing,
        df_invoice,
    )

    # Upload Word document to Active folder
    upload_file_to_onedrive(
        graph_token,
        user_email,
        active_folder,
        local_output_path,
        output_word_name,
    )

    # Convert to PDF using OneDrive
    onedrive_word_path = f"{active_folder}/{output_word_name}"
    convert_word_to_pdf_onedrive(
        graph_token,
        user_email,
        onedrive_word_path,
        local_pdf_path,
    )

    # Upload PDF to Active folder
    upload_file_to_onedrive(
        graph_token,
        user_email,
        active_folder,
        local_pdf_path,
        pdf_output_name,
    )

    # Upload PDF to Dataverse
    upload_file_to_dataverse_file_column(
        dataverse_url,
        closing_table,
        closing_row_id,
        dv_token,
        CLOSING_PDF_FILE_COLUMN,
        local_pdf_path,
        pdf_output_name,
    )

    # Clean up local files
    if os.path.exists(local_template_path):
        os.remove(local_template_path)
    if os.path.exists(local_output_path):
        os.remove(local_output_path)
    if os.path.exists(local_pdf_path):
        os.remove(local_pdf_path)

    print("\nProcess completed successfully")


if __name__ == "__main__":
    main()
